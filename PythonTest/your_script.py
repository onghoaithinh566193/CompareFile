from flask import Flask, request, render_template, send_file
import os
import pandas as pd
from openpyxl.styles import PatternFill
from docx import Document

app = Flask(__name__)

def compare_docs_and_log_changes(doc1_path, doc2_path, excel_path):
    # Mở hai file Word
    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)

    changes = []

    # So sánh đoạn văn
    max_paragraphs = max(len(doc1.paragraphs), len(doc2.paragraphs))
    for i in range(max_paragraphs):
        text1 = doc1.paragraphs[i].text if i < len(doc1.paragraphs) else ""
        text2 = doc2.paragraphs[i].text if i < len(doc2.paragraphs) else ""
        if text1 != text2:
            changes.append({
                'Nội dung cũ': text1,
                'Nội dung mới': text2
            })

    # So sánh bảng
    max_tables = max(len(doc1.tables), len(doc2.tables))
    for t_idx in range(max_tables):
        t1 = doc1.tables[t_idx] if t_idx < len(doc1.tables) else None
        t2 = doc2.tables[t_idx] if t_idx < len(doc2.tables) else None
        if t1 and t2:
            max_rows = max(len(t1.rows), len(t2.rows))
            max_cols = max(len(t1.columns), len(t2.columns))
            for r in range(max_rows):
                for c in range(max_cols):
                    cell_text1 = t1.cell(r, c).text if r < len(t1.rows) and c < len(t1.columns) else ""
                    cell_text2 = t2.cell(r, c).text if r < len(t2.rows) and c < len(t2.columns) else ""
                    if cell_text1 != cell_text2:
                        changes.append({
                            'Nội dung cũ': cell_text1,
                            'Nội dung mới': cell_text2
                        })

    # Ghi ra Excel
    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        df = pd.DataFrame(changes)
        df.to_excel(writer, index=False, sheet_name='Thay đổi')
        wb = writer.book
        ws = wb['Thay đổi']

        fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')

        # Tô màu cột "Nội dung cũ"
        header_row = ws[1]
        col_index_old_content = None
        for cell in header_row:
            if cell.value == 'Nội dung cũ':
                col_index_old_content = cell.column
                break

        if col_index_old_content:
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
                cell = row[col_index_old_content - 1]
                if cell.value:
                    cell.fill = fill

        # Căn chỉnh cột
        for col_idx, column_cells in enumerate(ws.columns, start=1):
            length = max(len(str(cell.value or '')) for cell in column_cells)
            col_letter = ws.cell(row=1, column=col_idx).coordinate[:1]
            ws.column_dimensions[col_letter].width = length + 2

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Lấy file upload
        file1 = request.files['doc1']
        file2 = request.files['doc2']
        # Lưu tạm các file
        path1 = os.path.join('uploads', file1.filename)
        path2 = os.path.join('uploads', file2.filename)
        file1.save(path1)
        file2.save(path2)
        # Đường dẫn xuất ra
        output_path = os.path.join('outputs', 'Ketqua.xlsx')
        # Gọi hàm so sánh
        compare_docs_and_log_changes(path1, path2, output_path)
        # Gửi file về
        return send_file(output_path, as_attachment=True)
    return render_template('index.html')

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    os.makedirs('outputs', exist_ok=True)
    app.run(debug=True)
